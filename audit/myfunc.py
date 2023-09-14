from win32com import client
from PIL import Image
import os
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
import pandas as pd

def getfile(path:str, type_of_file=None) -> list:
    """查找目录下的所有文件

    Args:
        path (str): 绝对路径
        type_of_file (str, list[str]): 文件后缀名. Defaults to None.

    Returns:
        list: 文件的绝对路径
    """
    file_list = []  # 用来存储所有的file文件路径
    if type_of_file is None: # 不指定特定文档
        print('将查找所有的文件')
        for current_folder, list_folders, files in os.walk(path):
            for f in files:  # 用来遍历所有的文件，只取文件名，不取路径名
                file_list.append(current_folder + '\\' + f)
                
    if isinstance(type_of_file, str): # 单个文档类型
        for current_folder, list_folders, files in os.walk(path):
            for f in files:  # 用来遍历所有的文件，只取文件名，不取路径名
                if f.endswith(type_of_file):  # 判断file文档
                    file_list.append(current_folder + '\\' + f)  # 把路径添加到列表中
                    
    if isinstance(type_of_file, list): # 多个文档类型
        for i in type_of_file:
            for current_folder, list_folders, files in os.walk(path):
                for f in files:  # 用来遍历所有的文件，只取文件名，不取路径名
                    if f.endswith(i):  # 判断file文档
                        file_list.append(current_folder + '\\' + f)  # 把路径添加到列表中
                                        
    return file_list  # 返回这个word文档的路径

def jpg2pdf(jpgFile:str)->None:
    """jpg转pdf

    Args:
        jpgFile (str): 图片的绝对路径

    Returns:
        _type_: _description_
    """
    path,fileName = jpgFile.rsplit('\\',1)
    preName,postName = fileName.rsplit('.',1)
    img = Image.open(jpgFile)
    return img.save(path+"\\"+preName+'.pdf', "PDF", resolution=100.0, save_all=True)


# 转换doc为pdf 待优化
def doc2pdf(fn, path1)->None:
    word = client.Dispatch("Word.Application")  # 打开word应用程序
    doc = word.Documents.Open(fn)  # 打开word文件

    a = os.path.split(fn)  # 分离路径和文件
    b = os.path.splitext(a[-1])[0]  # 拿到文件名

    doc.SaveAs("{}\\{}.pdf".format(path1, b), 17)  # 另存为后缀为".pdf"的文件，其中参数17表示为pdf
    doc.Close()  # 关闭原来word文件
    word.Quit()


# 转换docx为pdf 待优化
def docx2pdf(fn, path1):
    word = client.Dispatch("Word.Application")  # 打开word应用程序
    doc = word.Documents.Open(fn)  # 打开word文件

    a = os.path.split(fn)  # 分离路径和文件
    b = os.path.splitext(a[-1])[0]  # 拿到文件名

    doc.SaveAs("{}\\{}.pdf".format(path1, b), 17)  # 另存为后缀为".pdf"的文件，其中参数17表示为pdf
    doc.Close()  # 关闭原来word文件
    word.Quit()

def docx_find_replace_text(doc:Document(), search_text:str, replace_text:str)->None:
    """替换doc正文中的文字

    Args:
        doc (Document): _description_
        search_text (_type_): _description_
        replace_text (_type_): _description_
    """
    for paragraph in doc.paragraphs:
        if search_text in paragraph.text:
            for run in paragraph.runs:
                if search_text in run.text:
                    run.text = run.text.replace(search_text, replace_text)


def docx_find_replace_header(doc, search_text, replace_text):
    """替换页眉中的文字

    Args:
        doc (_type_): _description_
        search_text (_type_): _description_
        replace_text (_type_): _description_
    """
    for section in doc.sections:
        header = section.header  # 获取节的页眉
        for paragraph in header.paragraphs:  # 遍历页眉的每个段落
            if search_text in paragraph.text:
                for run in paragraph.runs:
                    if search_text in run.text:
                        run.text = run.text.replace(search_text, replace_text)

# docx_find_replace_header(doc, '«c_id»', 'new text')

#确定样式
def set_cell_text_style(cell, pos): 
    """更改word文字样式

    Args:
        cell (_type_): _description_
        pos (_type_): _description_

    Raises:
        ValueError: _description_

    Returns:
        _type_: _description_
    """
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER # 垂直居中
    for para in cell.paragraphs:
        if pos.lower() == "left":
            para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        elif pos.lower() == "mid":
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        elif pos.lower() == "right":
            para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        else:
            raise ValueError("Invalid position value. Please use 'left', 'mid' or 'right'.")
        for run in para.runs:
            run.font.size = Pt(10.5) # 五号字体
            run.font.name = u"宋体"
            run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    return cell

def date2str(df:pd.DataFrame, date_list:list)->None: # 等待进一步完善
    """日期格式转字符串

    Args:
        df (pd.DataFrame): _description_
        date_list (list): _description_
    """
    for i in date_list:
        df[i] = pd.to_datetime(df[i])
        df[i] = df[i].dt.strftime('%Y-%m-%d').fillna('')
        print(i, '转换完成！')
        
